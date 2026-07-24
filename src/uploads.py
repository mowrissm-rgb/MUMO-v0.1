"""
MUMO — Upload readers (files → chemical entities)
Multi-Agent Drug Discovery & Development AI Platform
Author: Mowriss & Claude (research partner)

WHAT THIS DOES
--------------
Turns an uploaded file into something MUMO can act on, so a user can feed the
tool a structure or a compound list when a typed name won't resolve — the
"upload the ligand/target instead" and "decode this GC-MS report" cases.

Every reader returns ONE common shape, so nothing downstream has to know or
care what kind of file it came from:

    {"kind": "compound_list" | "ligand" | "target" | "empty" | "error",
     "compounds": [{"name": str, "smiles": str|None, "source": str}],
     "target":    {"name": str, "pdb_text": str} | None,
     "summary":   "<one plain-language line about what was read>",
     "needs_llm": bool,          # document text that still needs the LLM to
     "text":      str,           #   extract a compound list from (PDF/Word)
     "error":     str|None}

DESIGN NOTES
------------
* Pure and dependency-light: RDKit / pandas / python-docx are already in the
  stack; pypdf and openpyxl are pure-Python additions. No OCR, no native
  package — this repo has been taken down twice by heavy conda deps, so the
  bar for adding one is high and none of these clear it into "heavy".
* NEVER raises. A malformed or hostile file returns {"kind": "error", ...};
  the app must not crash on whatever a user drops in.
* Reads from BYTES, not a path — the caller hands us the uploaded buffer and
  we never persist the raw file (see the app layer). Size limiting is the
  caller's job (done at the Streamlit boundary), but we still guard against a
  file that parses to something absurd.
"""

import io
import os
import re

MAX_COMPOUNDS = 500          # a sane ceiling; a real GC-MS list is < ~100
STRUCTURE_EXT = {".sdf", ".mol", ".mol2", ".pdb"}
TABULAR_EXT = {".csv", ".tsv", ".xlsx", ".xls"}
DOC_EXT = {".pdf", ".docx"}
ALL_EXT = STRUCTURE_EXT | TABULAR_EXT | DOC_EXT


def _err(msg):
    return {"kind": "error", "compounds": [], "target": None,
            "summary": "", "needs_llm": False, "text": "", "error": msg}


def _ok(kind, compounds=None, target=None, summary="", needs_llm=False, text=""):
    return {"kind": kind, "compounds": compounds or [], "target": target,
            "summary": summary, "needs_llm": needs_llm, "text": text,
            "error": None}


def _ext(filename):
    return os.path.splitext(str(filename or "").lower())[1]


# ── RDKit structure files ──────────────────────────────────────────────────

def _mol_entry(mol, fallback_name, source):
    """One {name, smiles} from an RDKit mol, or None if it won't sanitize."""
    from rdkit import Chem
    if mol is None:
        return None
    try:
        smi = Chem.MolToSmiles(mol)
    except Exception:
        return None
    if not smi:
        return None
    name = ""
    for prop in ("_Name", "Name", "NAME", "ID", "Compound", "COMPOUND"):
        if mol.HasProp(prop):
            name = mol.GetProp(prop).strip()
            if name:
                break
    return {"name": name or fallback_name, "smiles": smi, "source": source}


def read_sdf(data, filename="upload.sdf"):
    """An SDF may hold ONE molecule or a whole library — treat it as a list and
    let the caller decide; a single-entry list is just a list of length one."""
    from rdkit import Chem
    text = _as_text(data)
    supplier = Chem.SDMolSupplier()
    try:
        supplier.SetData(text, sanitize=True, removeHs=False)
    except Exception as e:
        return _err(f"Couldn't parse the SDF: {e}")
    base = os.path.splitext(os.path.basename(filename))[0]
    compounds = []
    for i, mol in enumerate(supplier):
        entry = _mol_entry(mol, f"{base}_{i + 1}", f"SDF ({filename})")
        if entry:
            compounds.append(entry)
        if len(compounds) >= MAX_COMPOUNDS:
            break
    if not compounds:
        return _err("No readable molecules in that SDF.")
    kind = "compound_list" if len(compounds) > 1 else "ligand"
    return _ok(kind, compounds=compounds,
               summary=f"Read {len(compounds)} molecule(s) from {filename}.")


def read_molblock(data, filename="upload.mol"):
    """A single MOL/SDF-molblock structure → one ligand."""
    from rdkit import Chem
    text = _as_text(data)
    mol = Chem.MolFromMolBlock(text, sanitize=True, removeHs=False)
    if mol is None:
        mol = Chem.MolFromMolBlock(text, sanitize=False, removeHs=False)
    entry = _mol_entry(mol, os.path.splitext(os.path.basename(filename))[0],
                       f"MOL ({filename})")
    if not entry:
        return _err("Couldn't read a molecule from that MOL file.")
    return _ok("ligand", compounds=[entry],
               summary=f"Read one molecule from {filename}.")


def read_mol2(data, filename="upload.mol2"):
    from rdkit import Chem
    text = _as_text(data)
    mol = Chem.MolFromMol2Block(text, sanitize=True, removeHs=False)
    if mol is None:
        mol = Chem.MolFromMol2Block(text, sanitize=False, removeHs=False)
    entry = _mol_entry(mol, os.path.splitext(os.path.basename(filename))[0],
                       f"MOL2 ({filename})")
    if not entry:
        return _err("Couldn't read a molecule from that MOL2 file.")
    return _ok("ligand", compounds=[entry],
               summary=f"Read one molecule from {filename}.")


# ── PDB: protein target vs small-molecule ligand ───────────────────────────

_AMINO3 = {"ALA", "ARG", "ASN", "ASP", "CYS", "GLN", "GLU", "GLY", "HIS", "ILE",
           "LEU", "LYS", "MET", "PHE", "PRO", "SER", "THR", "TRP", "TYR", "VAL"}


def read_pdb(data, filename="upload.pdb"):
    """Decide whether a PDB is a PROTEIN (a dockable target) or a small
    MOLECULE (a ligand), by counting standard amino-acid backbone residues.

    The distinction is what lets a single upload answer 'validate/dock this
    target' OR 'here is the ligand you couldn't name' — routing the wrong way
    would dock a protein as if it were a drug, or vice versa.
    """
    text = _as_text(data)
    if "ATOM" not in text and "HETATM" not in text:
        return _err("That PDB has no atom records.")
    amino_ca = 0
    hetatm = 0
    for line in text.splitlines():
        if line.startswith("ATOM") and line[12:16].strip() == "CA" \
                and line[17:20].strip() in _AMINO3:
            amino_ca += 1
        elif line.startswith("HETATM") and line[17:20].strip() not in ("HOH", "WAT"):
            hetatm += 1
    base = os.path.splitext(os.path.basename(filename))[0]
    if amino_ca >= 20:
        return _ok("target",
                   target={"name": base, "pdb_text": text},
                   summary=f"Read a protein structure ({amino_ca} residues) "
                           f"from {filename} — treating it as a target.")
    # Not enough protein — try to read it as a small molecule.
    from rdkit import Chem
    mol = Chem.MolFromPDBBlock(text, sanitize=True, removeHs=False)
    if mol is None:
        mol = Chem.MolFromPDBBlock(text, sanitize=False, removeHs=False)
    entry = _mol_entry(mol, base, f"PDB ({filename})")
    if entry:
        return _ok("ligand", compounds=[entry],
                   summary=f"Read a small molecule from {filename} — treating "
                           f"it as a ligand.")
    if amino_ca:      # some protein, just short (a peptide/fragment)
        return _ok("target", target={"name": base, "pdb_text": text},
                   summary=f"Read a short protein/peptide from {filename}.")
    return _err("Couldn't interpret that PDB as a protein or a molecule.")


# ── tabular: CSV / Excel ───────────────────────────────────────────────────

_SMILES_COLS = ("smiles", "smile", "canonical_smiles", "isomeric_smiles",
                "structure", "smi")
_INCHI_COLS = ("inchi", "std_inchi", "standard_inchi")
_NAME_COLS = ("name", "compound", "compound_name", "ligand", "molecule",
              "title", "id", "chemical", "compound name")
_TARGET_COLS = ("target", "protein", "gene", "receptor", "pdb", "pdb_id", "uniprot")


def _pick_col(columns, wanted):
    norm = {re.sub(r"[^a-z0-9]", "", str(c).lower()): c for c in columns}
    for w in wanted:
        key = re.sub(r"[^a-z0-9]", "", w)
        if key in norm:
            return norm[key]
    return None


def read_tabular(data, filename="upload.csv"):
    """CSV/Excel → compounds (SMILES + name columns) or targets.

    Auto-detects a SMILES/InChI column and a name column so a plain export
    ('Compound, SMILES, ...') just works. If there is no structure column but
    a name column, the names are returned unresolved for the normal
    name→structure path to handle downstream.
    """
    import pandas as pd
    ext = _ext(filename)
    try:
        if ext in (".xlsx", ".xls"):
            df = pd.read_excel(io.BytesIO(_as_bytes(data)))
        else:
            df = pd.read_csv(io.BytesIO(_as_bytes(data)),
                             sep=_sniff_delimiter(_as_text(data), ext))
    except Exception as e:
        return _err(f"Couldn't read that table: {e}")
    if df.empty:
        return _err("That table is empty.")

    smi_col = _pick_col(df.columns, _SMILES_COLS)
    inchi_col = None if smi_col else _pick_col(df.columns, _INCHI_COLS)
    name_col = _pick_col(df.columns, _NAME_COLS)
    target_col = _pick_col(df.columns, _TARGET_COLS)

    compounds = []
    for _, row in df.iterrows():
        if len(compounds) >= MAX_COMPOUNDS:
            break
        smi = None
        if smi_col is not None:
            smi = _clean_cell(row.get(smi_col)) or None   # empty cell -> None
        elif inchi_col is not None:
            smi = _inchi_to_smiles(_clean_cell(row.get(inchi_col)))
        name = _clean_cell(row.get(name_col)) if name_col is not None else ""
        if not smi and not name:
            continue
        compounds.append({"name": name or "(unnamed)", "smiles": smi,
                          "source": f"table ({filename})"})

    if compounds:
        kind = "compound_list" if len(compounds) > 1 else "ligand"
        with_struct = sum(1 for c in compounds if c["smiles"])
        return _ok(kind, compounds=compounds,
                   summary=(f"Read {len(compounds)} compound(s) from {filename} "
                            f"({with_struct} with a structure column)."))

    # No compound columns — maybe it is a list of targets.
    if target_col is not None:
        names = [_clean_cell(row.get(target_col)) for _, row in df.iterrows()]
        names = [n for n in names if n][:MAX_COMPOUNDS]
        if names:
            return _ok("target_list",
                       compounds=[{"name": n, "smiles": None, "source": "table"}
                                  for n in names],
                       summary=f"Read {len(names)} target name(s) from {filename}.")

    return _err(f"Couldn't find a compound, SMILES, or target column in "
                f"{filename}. Columns were: {', '.join(map(str, df.columns))}.")


def _sniff_delimiter(text, ext):
    """Pick a CSV/TSV delimiter from a fixed candidate set by counting it in
    the header row. csv.Sniffer guesses among ALL characters and picks a
    letter on a single-column file ('target' -> delimiter 't'), which is a
    known failure; restricting to real delimiters avoids it. A single-column
    file simply has zero of every candidate, and comma is the safe default.
    """
    if ext == ".tsv":
        return "\t"
    header = (text or "").splitlines()[0] if text else ""
    best, best_n = ",", 0
    for cand in (",", ";", "\t", "|"):
        n = header.count(cand)
        if n > best_n:
            best, best_n = cand, n
    return best


def _inchi_to_smiles(inchi):
    if not inchi:
        return None
    try:
        from rdkit import Chem
        m = Chem.MolFromInchi(inchi)
        return Chem.MolToSmiles(m) if m else None
    except Exception:
        return None


# ── documents: Word / PDF → text (LLM extracts the compound list) ──────────

def read_docx(data, filename="upload.docx"):
    try:
        from docx import Document
        doc = Document(io.BytesIO(_as_bytes(data)))
    except Exception as e:
        return _err(f"Couldn't open that Word file: {e}")
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:               # GC-MS results are often tables
        for trow in table.rows:
            cells = [c.text.strip() for c in trow.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        return _err("That Word file had no readable text.")
    return _ok("document", needs_llm=True, text=text[:60000],
               summary=f"Read text from {filename}; extracting the compound list…")


def read_pdf(data, filename="upload.pdf"):
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(_as_bytes(data)))
    except Exception as e:
        return _err(f"Couldn't open that PDF: {e}")
    pages = []
    for page in reader.pages[:40]:         # cap: a report, not a book
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pass
    text = "\n".join(pages).strip()
    if not text:
        return _err("No selectable text in that PDF — it may be a scanned "
                    "image, which isn't supported yet. Upload the data as CSV/"
                    "Excel, or paste the compound names.")
    return _ok("document", needs_llm=True, text=text[:60000],
               summary=f"Read text from {filename}; extracting the compound list…")


_GCMS_SYSTEM = (
    "You extract a clean compound list from the raw text of a chemistry "
    "document — usually a GC-MS / phytochemical analysis report, sometimes a "
    "plain list. The text was pulled from a PDF or Word file and may be messy: "
    "broken table rows, headers, page numbers, retention times, peak areas.\n\n"
    "Return ONLY the chemical COMPOUNDS that are actual analytes — the molecules "
    "the study detected. IGNORE everything else: column headers, instrument "
    "settings, author names, retention-time/area/percentage numbers on their "
    "own, references, page furniture.\n\n"
    "For each compound give its name EXACTLY as written in the document "
    "(preserve CAS-style names like 'Cholan-24-oic acid, 3,12-dioxo-, (5.beta.)-' "
    "verbatim — do not tidy or rename them), and its area/percentage if one is "
    "clearly associated with it, else null.\n\n"
    "Reply with ONLY a JSON array, nothing else:\n"
    '[{"name": "<compound name as written>", "percent": <number|null>}, ...]\n'
    "If the text contains no identifiable compounds, return []."
)


def extract_compounds_from_text(text, llm, source="document"):
    """Pull a compound list out of document text using the LLM.

    Kept here (not in the chat layer) so the GC-MS decode is unit-testable with
    a mocked llm. `llm` is any object with .chat(system, user) -> str, matching
    what the rest of MUMO uses. Returns the same {"compounds": [...]} shape as
    the structure readers, so a decoded PDF and an uploaded SDF are
    interchangeable downstream. Never raises.
    """
    import json
    if not text or llm is None:
        return _ok("compound_list", compounds=[],
                   summary="No text or no language model available to read it.")
    try:
        raw = llm.chat(_GCMS_SYSTEM, text[:60000], temperature=0.0, max_tokens=2000)
    except Exception as e:
        return _err(f"Couldn't read the document's compounds: {e}")
    # the model is asked for a bare array, but be forgiving about stray prose
    m = re.search(r"\[.*\]", raw or "", re.DOTALL)
    if not m:
        return _ok("compound_list", compounds=[],
                   summary="No compounds could be identified in that document.")
    try:
        items = json.loads(m.group(0))
    except Exception:
        return _ok("compound_list", compounds=[],
                   summary="The document's compound list couldn't be parsed.")
    compounds, seen = [], set()
    for it in items if isinstance(items, list) else []:
        name = _clean_cell((it or {}).get("name") if isinstance(it, dict) else it)
        if not name or name.lower() in seen:
            continue
        seen.add(name.lower())
        pct = (it.get("percent") if isinstance(it, dict) else None)
        compounds.append({"name": name, "smiles": None, "source": source,
                          "percent": pct})
        if len(compounds) >= MAX_COMPOUNDS:
            break
    return _ok("compound_list", compounds=compounds,
               summary=f"Identified {len(compounds)} compound(s) in the document.")


# ── dispatch ───────────────────────────────────────────────────────────────

def read_upload(filename, data):
    """Route an uploaded file to the right reader by extension. Never raises."""
    ext = _ext(filename)
    try:
        if ext == ".sdf":
            return read_sdf(data, filename)
        if ext == ".mol":
            return read_molblock(data, filename)
        if ext == ".mol2":
            return read_mol2(data, filename)
        if ext == ".pdb":
            return read_pdb(data, filename)
        if ext in TABULAR_EXT:
            return read_tabular(data, filename)
        if ext == ".docx":
            return read_docx(data, filename)
        if ext == ".pdf":
            return read_pdf(data, filename)
    except Exception as e:
        return _err(f"Couldn't read {filename}: {type(e).__name__}: {e}")
    return _err(f"Unsupported file type '{ext or '?'}'. Supported: "
                f"{', '.join(sorted(ALL_EXT))}.")


# ── small helpers ──────────────────────────────────────────────────────────

def _as_bytes(data):
    if isinstance(data, bytes):
        return data
    if isinstance(data, str):
        return data.encode("utf-8", "replace")
    if hasattr(data, "read"):
        return data.read()
    return bytes(data)


def _as_text(data):
    b = _as_bytes(data)
    for enc in ("utf-8", "latin-1"):
        try:
            return b.decode(enc)
        except Exception:
            continue
    return b.decode("utf-8", "replace")


def _clean_cell(v):
    if v is None:
        return ""
    s = str(v).strip()
    return "" if s.lower() in ("nan", "none", "") else s
