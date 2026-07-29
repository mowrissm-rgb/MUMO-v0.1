"""
MUMO — Report Writer
Multi-Agent Drug Discovery & Development AI Platform

Turns a pipeline's results (docking / STRING / ADMET) into a downloadable
.docx: tables, write-ups, and static images. 2D/3D/network graphics are
rasterized through a headless Chromium (Playwright) so the exported picture
matches exactly what's shown live in the app (same SVG, same 3Dmol.js pose).

Heavy imports (docx, playwright, pandas, brain) are lazy — this module only
gets touched when a user actually clicks "Generate report".
"""

import io
import re


def normalize_svg_viewbox(svg):
    """Add a viewBox derived from px width/height if missing (STRING's raw SVG
    has none), so the SVG scales instead of clipping when rasterized."""
    m = re.search(r"<svg[^>]*?>", svg)
    if not m:
        return svg
    tag = m.group(0)
    if "viewbox" in tag.lower():
        return svg
    wm = re.search(r'width=["\']([\d.]+)', tag)
    hm = re.search(r'height=["\']([\d.]+)', tag)
    if not (wm and hm):
        return svg
    newtag = re.sub(r"<svg", f'<svg viewBox="0 0 {wm.group(1)} {hm.group(1)}"', tag, count=1)
    return svg.replace(tag, newtag, 1)


def new_browser():
    """Launch one headless Chromium instance to reuse across every screenshot
    in a report (much faster than a fresh browser per image).

    The 3D pose view needs WebGL (3Dmol.js/three.js) — default headless Chromium
    creates a WebGL context but then immediately loses it (no GPU in a headless
    container), which silently renders nothing. These flags force software
    rendering (SwiftShader) so WebGL actually stays alive and draws."""
    from playwright.sync_api import sync_playwright
    pw = sync_playwright().start()
    browser = pw.chromium.launch(args=[
        "--use-gl=angle", "--use-angle=swiftshader", "--enable-webgl",
        "--ignore-gpu-blocklist", "--enable-unsafe-swiftshader", "--disable-gpu-sandbox",
    ])
    return pw, browser


def svg_to_png(svg, browser, width=760, height=560, pad=16):
    """Rasterize an SVG string to PNG bytes via headless Chromium."""
    svg = normalize_svg_viewbox(svg)
    html = (f'<html><body style="margin:0;background:#fff;">'
            f'<div id="c" style="display:inline-block;padding:{pad}px;background:#fff;">'
            f'{svg}</div></body></html>')
    page = browser.new_page(viewport={"width": width, "height": height})
    try:
        page.set_content(html)
        page.wait_for_timeout(150)
        return page.locator("#c").screenshot()
    finally:
        page.close()


def png_from_3d(complex_pdb_path, ia, browser, options=None, width=900, height=560):
    """Load MUMO's own 3Dmol viewer HTML in headless Chromium and screenshot the
    rendered pose — the same view shown in-app, captured as a static image.
    Raises on failure (e.g. the CDN 3Dmol.js load times out) so the caller can
    report the real reason instead of silently omitting the image."""
    from viz import render_complex_html
    html = render_complex_html(complex_pdb_path, ia, options=options, width=width, height=height)
    page = browser.new_page(viewport={"width": width, "height": height + 20})
    try:
        page.set_content(f"<html><body style='margin:0;'>{html}</body></html>")
        page.wait_for_function("window.__mumoReady === true", timeout=20000)
        page.wait_for_timeout(200)  # let the final zoom/render settle
        return page.locator("#mumoview").screenshot()
    finally:
        page.close()


# ─────────────────────────────────────────────────────────────────────────────
# docx assembly helpers
# ─────────────────────────────────────────────────────────────────────────────

def _add_kv_table(doc, pairs):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for k, v in pairs:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = str(k), str(v)
    return table


def _add_df_table(doc, df):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = str(row[col])
    return table


def _add_bold_runs(paragraph, text):
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part:
            paragraph.add_run(part)


def _add_markdown(doc, text):
    """Small markdown -> docx converter: headings, bullet lines, **bold** spans,
    plain paragraphs. Enough to render the LLM's beginner-narrative reports."""
    if not text:
        return
    for line in text.split("\n"):
        line = line.rstrip()
        if not line:
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith(("- ", "* ")):
            _add_bold_runs(doc.add_paragraph(style="List Bullet"), line[2:])
        else:
            _add_bold_runs(doc.add_paragraph(), line)


# ─────────────────────────────────────────────────────────────────────────────
# per-pipeline report builders
# ─────────────────────────────────────────────────────────────────────────────

def build_docking_docx(r, llm=None):
    """Full docking report, organised BY TARGET.

    The previous layout put every ligand-target pair into one flat table and
    one flat chart, then gave all 60 pairs their own section. On a 15-ligand
    4-target screen that is a 61-row table, a 60-bar chart and 116 images —
    nothing in it can actually be read, and the file ran to 10 MB.

    This groups by target instead. Each target gets its own ranked table, its
    own three figures and ONE pose (its best ligand), so every number sits
    next to the protein it belongs to. Tables and figures are numbered
    sequentially and cross-referenced, so a reader can follow a residue list
    from a table row to the pose that produced it.

    Section order is deliberate and was specified: validate the structure
    first, then the ligands, and only then present affinities that depend on
    both being sound.
    """
    from docx import Document
    from docx.shared import Inches, Pt
    from brain import write_report
    from viz import find_entry

    rdf, viz, meta = r["rdf"], r.get("viz", {}), r.get("meta", {})
    doc = Document()
    doc.add_heading("MUMO Docking Report", level=0)

    # ── numbering ──────────────────────────────────────────────────────────
    # One counter per kind, shared across the whole document, so captions can
    # refer to each other ("matches the top row of Table 1") the way a paper
    # does. Without this the reader has no way to tie a pose to its row.
    counts = {"Table": 0, "Figure": 0}

    def caption(kind, text):
        counts[kind] += 1
        p = doc.add_paragraph()
        run = p.add_run(f"{kind} {counts[kind]}. {text}")
        run.italic = True
        run.font.size = Pt(9)
        return counts[kind]

    def picture(png, width=6.0):
        doc.add_picture(io.BytesIO(png), width=Inches(width))

    # ── rows, grouped by target ────────────────────────────────────────────
    rows = rdf.reset_index().rename(columns={"index": "Rank"}).to_dict(orient="records")
    ok_rows = [x for x in rows if str(x.get("Best affinity (kcal/mol)")) != "FAILED"]
    failed_rows = [x for x in rows if str(x.get("Best affinity (kcal/mol)")) == "FAILED"]

    def _aff(x):
        try:
            return float(x.get("Best affinity (kcal/mol)"))
        except (TypeError, ValueError):
            return 0.0

    has_targets = any(x.get("Target") for x in rows)
    if has_targets:
        targets = []
        for x in rows:
            t = x.get("Target")
            if t and t not in targets:
                targets.append(t)
    else:
        targets = [meta.get("gene", "target")]

    def rows_for(t):
        sub = [x for x in ok_rows if (x.get("Target") == t)] if has_targets else list(ok_rows)
        return sorted(sub, key=_aff)          # most negative (best) first

    # ── headless browser, shared ───────────────────────────────────────────
    bh = {"pw": None, "browser": None, "err": None}
    try:
        bh["pw"], bh["browser"] = new_browser()
    except Exception as e:
        bh["err"] = f"{type(e).__name__}: {e}"

    def _shot(kind, *args):
        """Screenshot with one restart if the browser died mid-batch."""
        for attempt in (0, 1):
            if bh["browser"] is None:
                return None, bh["err"] or "headless browser unavailable"
            try:
                if kind == "2d":
                    return svg_to_png(args[0], bh["browser"]), None
                return png_from_3d(args[0], args[1], bh["browser"]), None
            except Exception as e:
                bh["err"] = f"{type(e).__name__}: {e}"
                if attempt == 0:
                    try:
                        bh["browser"].close(); bh["pw"].stop()
                    except Exception:
                        pass
                    try:
                        bh["pw"], bh["browser"] = new_browser()
                    except Exception as e2:
                        bh["browser"] = None
                        return None, f"{type(e2).__name__}: {e2}"
        return None, bh["err"]

    try:
        # ══ SUMMARY ════════════════════════════════════════════════════════
        doc.add_heading("Summary", level=1)
        n_lig = len({x.get("Ligand") for x in rows})
        if len(targets) > 1:
            doc.add_paragraph(
                f"This run docked {n_lig} small-molecule ligands against each of "
                f"{len(targets)} targets — {', '.join(targets)} — for a total of "
                f"{len(ok_rows)} scored poses.")
        else:
            doc.add_paragraph(
                f"This run docked {n_lig} small-molecule ligand(s) against "
                f"{targets[0]}, producing {len(ok_rows)} scored poses.")

        best_bits = []
        for t in targets:
            sub = rows_for(t)
            if sub:
                b = sub[0]
                best_bits.append(f"{t} → {b['Ligand']} ({_aff(b):.3f} kcal/mol)")
        if best_bits:
            doc.add_paragraph("Per-target best hit: " + "; ".join(best_bits) + ".")

        method = []
        if meta.get("exhaustiveness"):
            method.append(f"exhaustiveness {meta['exhaustiveness']}")
        if meta.get("replicas"):
            method.append(f"{meta['replicas']} replica(s)")
        if method:
            doc.add_paragraph("Method: AutoDock Vina · " + " · ".join(method) + ".")
        if failed_rows:
            doc.add_paragraph(
                f"{len(failed_rows)} ligand(s) could not be docked and are "
                f"excluded from the tables below: "
                + ", ".join(sorted({x['Ligand'] for x in failed_rows})) + ".")

        # ══ 1. STRUCTURE VALIDATION ════════════════════════════════════════
        # First, because every affinity below is conditional on the receptor
        # geometry being sound. A distorted pocket invalidates the numbers.
        doc.add_heading("Structure validation", level=1)
        per_target = meta.get("per_target") or {}
        wrote_any = False
        for t in targets:
            tm = per_target.get(t) or ({} if len(targets) > 1 else meta)
            rama = (tm or {}).get("ramachandran") or (
                meta.get("ramachandran") if len(targets) == 1 else None)
            pocket_bits = []
            if (tm or {}).get("pocket"):
                pocket_bits.append(tm["pocket"])
            tval = (tm or {}).get("validation")
            if tval:
                pocket_bits.append(
                    f"native redock RMSD {tval['rmsd']} Å "
                    f"({'validated' if tval['passed'] else 'above 2 Å'})")

            if not (rama or pocket_bits):
                continue
            wrote_any = True
            doc.add_heading(t, level=2)

            if rama:
                try:
                    import ramachandran as _ram
                    doc.add_paragraph(_ram.verdict(rama))
                    svg = _ram.plot_svg(rama, title=f"Ramachandran plot — {t}")
                    if svg:
                        png, err = _shot("2d", svg)
                        if png:
                            picture(png, 5.0)
                            caption("Figure",
                                    f"Ramachandran plot for {t}. Backbone torsion angles "
                                    f"scored against the MolProbity Top8000 reference "
                                    f"distributions.")
                        else:
                            doc.add_paragraph(f"(Ramachandran plot unavailable: {err})")
                    outs = rama.get("outliers") or []
                    if outs:
                        shown = ", ".join(outs[:20]) + (" …" if len(outs) > 20 else "")
                        doc.add_paragraph(f"Outlier residues ({len(outs)}): {shown}")
                        doc.add_paragraph(
                            "Check whether any outlier lines the binding site — a distorted "
                            "residue in the pocket undermines the poses below far more than "
                            "one on a surface loop.")
                except Exception as e:
                    doc.add_paragraph(f"(Structure validation unavailable: "
                                      f"{type(e).__name__}: {e})")

            if pocket_bits:
                doc.add_paragraph("Docking site: " + " · ".join(pocket_bits) + ".")
        if not wrote_any:
            doc.add_paragraph("No structure-validation data was recorded for this run.")

        # ══ 2. LIGAND VALIDATION (HOMO/LUMO) ═══════════════════════════════
        # Second, because the electronic character of a ligand is a property of
        # the molecule alone — it does not depend on which target it was docked
        # into, so it belongs before the target-by-target results.
        doc.add_heading("Ligand validation — frontier molecular orbitals", level=1)
        qm_rows, qm_figs = [], []
        try:
            from agents.qm_analyst import orbitals, xtb_available
            if not xtb_available():
                doc.add_paragraph(
                    "The quantum-chemistry engine (xtb) is not available in this "
                    "build, so HOMO/LUMO properties were not computed.")
            else:
                seen = {}
                for x in ok_rows:
                    lab, smi = x.get("Ligand"), x.get("SMILES")
                    if lab and smi and lab not in seen:
                        seen[lab] = smi
                doc.add_paragraph(
                    f"Electronic properties of the {len(seen)} unique ligand(s) in this "
                    f"screen, computed with GFN2-xTB. The HOMO–LUMO gap indicates how "
                    f"readily a molecule takes part in electron transfer: a smaller gap "
                    f"means a more reactive, more easily polarised molecule.")
                for lab, smi in seen.items():
                    q = orbitals(smi)
                    if q.get("_error"):
                        qm_rows.append({"Ligand": lab, "HOMO (eV)": "—",
                                        "LUMO (eV)": "—", "Gap (eV)": "—",
                                        "Note": q["_error"][:60]})
                        continue
                    qm_rows.append({
                        "Ligand": lab,
                        "HOMO (eV)": f"{q['homo_ev']:.2f}",
                        "LUMO (eV)": f"{q['lumo_ev']:.2f}",
                        "Gap (eV)": f"{q['gap_ev']:.2f}",
                        "Note": "",
                    })
                    qm_figs.append((lab, q))
                if qm_rows:
                    import pandas as _pd
                    caption("Table", "Frontier orbital energies for every ligand in the "
                                     "screen (GFN2-xTB).")
                    _add_df_table(doc, _pd.DataFrame(qm_rows))
                for lab, q in qm_figs:
                    try:
                        import viz_string as _vs
                        svg = _vs.orbital_svg(q, dark=False,
                                              title=f"Frontier orbitals — {lab}")
                        if svg:
                            png, err = _shot("2d", svg)
                            if png:
                                picture(png, 4.6)
                                caption("Figure",
                                        f"Orbital energy diagram for {lab}: HOMO "
                                        f"{q['homo_ev']:.2f} eV, LUMO {q['lumo_ev']:.2f} eV, "
                                        f"gap {q['gap_ev']:.2f} eV.")
                    except Exception:
                        pass
        except Exception as e:
            doc.add_paragraph(f"(Ligand validation unavailable: {type(e).__name__}: {e})")

        # ══ 3. NARRATIVES (fetched concurrently, used below) ════════════════
        def _sp(v):
            return [x for x in str(v).split("; ") if x and x != "-"]

        def _narrative_job(row):
            label = row["Ligand"]
            row_target = row.get("Target") or meta.get("gene")
            # reliability_by is rebuilt PER TARGET and keyed by plain ligand
            # label, so the top-level dict only ever holds the FIRST target's.
            # Reading it for a row from another target attaches the wrong
            # target's reasoning — wrong data, not merely missing.
            if row.get("Target"):
                src = (meta.get("per_target", {}).get(row_target, {})
                       .get("reliability_by") or {})
            else:
                src = meta.get("reliability_by") or {}
            rel = src.get(label, {})
            try:
                return write_report({
                    "target": row_target, "ligand": label,
                    "affinity": _aff(row),
                    "estimated_ki": row.get("Est. Ki"),
                    "ligand_efficiency": row.get("Ligand efficiency"),
                    "reliability": row.get("Reliability"),
                    "reliability_reason": rel.get("reason"),
                    "total_interactions": row.get("Total interactions"),
                    "n_hbonds": int(row.get("H-bonds", 0) or 0),
                    "hbond_residues": _sp(row.get("H-bond residues", "")),
                    "n_hydrophobic": int(row.get("Hydrophobic", 0) or 0),
                    "interacting_residues": _sp(row.get("All interacting residues", "")),
                }, llm, r.get("tier", "Standard"))
            except Exception as e:
                return f"(Interpretation unavailable: {type(e).__name__}: {e})"

        from concurrent.futures import ThreadPoolExecutor
        top_rows = [rows_for(t)[0] for t in targets if rows_for(t)]
        writeups = {}
        if top_rows:
            with ThreadPoolExecutor(max_workers=min(6, len(top_rows))) as pool:
                futures = {id(x): pool.submit(_narrative_job, x) for x in top_rows}
                for x in top_rows:
                    writeups[id(x)] = futures[id(x)].result()

        # ══ 4. DOCKING RESULTS, PER TARGET ═════════════════════════════════
        doc.add_heading("Docking results", level=1)
        import charts
        table_no_for = {}

        for t in targets:
            sub = rows_for(t)
            if not sub:
                continue
            doc.add_heading(t, level=2)
            best = sub[0]
            ki = f" (Ki {best.get('Est. Ki')})" if best.get("Est. Ki") else ""
            doc.add_paragraph(
                f"Top ligand: {best['Ligand']} — {_aff(best):.3f} kcal/mol{ki}.")

            # ranked table
            doc.add_heading("Ranked ligand table", level=3)
            table_no = caption(
                "Table",
                f"Ranked docking results for all {len(sub)} ligand(s) screened against "
                f"{t}, best to worst binding affinity. The top row's residue list "
                f"matches the pose figure for this target.")
            table_no_for[t] = table_no
            import pandas as _pd
            tab = _pd.DataFrame([{
                "Ligand": x.get("Ligand"),
                "Affinity (kcal/mol)": f"{_aff(x):.3f}",
                "Est. Ki": x.get("Est. Ki", "—"),
                "Ligand eff.": x.get("Ligand efficiency", "—"),
                "H-bonds": x.get("H-bonds", 0),
                "Key interacting residues": x.get("All interacting residues", "—"),
            } for x in sub])
            _add_df_table(doc, tab)

            # figures, scoped to THIS target only
            for heading, fn, cap in (
                ("Binding affinity chart", charts.affinity_chart_svg,
                 f"Binding affinity by ligand for {t} (colour scale = affinity strength)."),
                ("Residue contact frequency", charts.residue_frequency_svg,
                 f"Frequency with which each binding-site residue in {t} is contacted "
                 f"across the docked ligands."),
                ("Ligand × residue heatmap", charts.contact_heatmap_svg,
                 f"Contact heatmap for {t}: rows = ligands (ranked by affinity), "
                 f"columns = residues (ranked by contact frequency)."),
            ):
                try:
                    svg = (fn(sub) if fn is charts.affinity_chart_svg
                           else fn(sub, target_name=t))
                    if not svg:
                        continue
                    doc.add_heading(heading, level=3)
                    png, err = _shot("2d", svg)
                    if png:
                        picture(png, 6.2)
                        caption("Figure", cap)
                    else:
                        doc.add_paragraph(f"({heading} unavailable: {err})")
                except Exception as e:
                    doc.add_paragraph(f"({heading} unavailable: {type(e).__name__}: {e})")

            # ONE pose — the best ligand for this target
            doc.add_heading(f"Top pose: {best['Ligand']}", level=3)
            _add_markdown(doc, writeups.get(id(best)) or "")
            entry = find_entry(viz, best["Ligand"], best.get("Target"))
            if not entry:
                doc.add_paragraph("(No pose data available for this ligand.)")
            else:
                shown = False
                png3d, err3 = _shot("3d", entry["complex"], entry["ia"])
                if png3d:
                    picture(png3d, 5.5)
                    shown = True
                svg2d = entry["ia"].get("svg_2d")
                if svg2d:
                    png2d, err2 = _shot("2d", svg2d)
                    if png2d:
                        picture(png2d, 5.5)
                        shown = True
                if shown:
                    caption("Figure",
                            f"3D binding pose and 2D interaction diagram of "
                            f"{best['Ligand']} in the {t} pocket. Residues shown here "
                            f"match the top row of Table {table_no_for.get(t, '?')}.")
                else:
                    doc.add_paragraph(f"(Pose images unavailable: {err3})")

        # ══ 5. RESULTS + 6. DISCUSSION ═════════════════════════════════════
        _add_results_discussion(doc, targets, rows_for, _aff, meta, llm, qm_rows)

    finally:
        if bh["browser"]:
            bh["browser"].close()
        if bh["pw"]:
            bh["pw"].stop()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_RESULTS_SYSTEM = (
    "You are writing the Results and Discussion sections of a molecular-docking "
    "report for a pharmacy researcher. Be precise, sober and academic. Use the "
    "standard AutoDock Vina interpretation bands: more negative than -7.0 kcal/mol "
    "is generally considered strong, -5.0 to -7.0 moderate, weaker than -5.0 poor. "
    "Never claim a result the numbers do not support, and say plainly when a screen "
    "produced no strong binder. Cover, under a '## Results' heading: what was docked "
    "against what, the per-target best hits with their affinities, and how the set "
    "distributes across those bands. Then under a '## Discussion' heading cover: the "
    "evaluation basis, why these targets matter, interpretation against the bands, "
    "whether any ligand is active across multiple targets, limitations (single "
    "replica, exhaustiveness setting, no solvent, docking scores are not free "
    "energies), and a short conclusion. Use markdown headings and short paragraphs."
)


def _add_results_discussion(doc, targets, rows_for, aff, meta, llm, qm_rows):
    """The written sections. Falls back to a factual template without an LLM —
    a report that simply stops after the tables is harder to act on."""
    facts = []
    for t in targets:
        sub = rows_for(t)
        if not sub:
            continue
        best = sub[0]
        facts.append(
            f"{t}: {len(sub)} ligands docked; best {best['Ligand']} at "
            f"{aff(best):.3f} kcal/mol (Ki {best.get('Est. Ki', 'n/a')}, "
            f"ligand efficiency {best.get('Ligand efficiency', 'n/a')}, "
            f"{best.get('H-bonds', 0)} H-bonds); "
            f"weakest {aff(sub[-1]):.3f} kcal/mol.")
    if qm_rows:
        gaps = [x for x in qm_rows if x.get("Gap (eV)") not in ("—", "", None)]
        if gaps:
            facts.append("HOMO-LUMO gaps (eV): " + ", ".join(
                f"{x['Ligand']} {x['Gap (eV)']}" for x in gaps[:12]))
    settings = []
    if meta.get("exhaustiveness"):
        settings.append(f"exhaustiveness {meta['exhaustiveness']}")
    if meta.get("replicas"):
        settings.append(f"{meta['replicas']} replica(s)")
    if settings:
        facts.append("Settings: " + ", ".join(settings) + ".")

    text = ""
    if llm:
        try:
            text = llm.chat(_RESULTS_SYSTEM,
                            "Run facts:\n" + "\n".join(facts) +
                            "\n\nWrite the Results and Discussion sections.",
                            temperature=0.3, max_tokens=1400)
        except Exception:
            text = ""

    if text.strip():
        _add_markdown(doc, text)
        return

    doc.add_heading("Results", level=1)
    for f in facts:
        doc.add_paragraph(f)
    doc.add_heading("Discussion", level=1)
    doc.add_paragraph(
        "Evaluation basis. Binding affinity more negative than -7.0 kcal/mol is "
        "generally considered strong for AutoDock Vina, -5.0 to -7.0 moderate, and "
        "weaker than -5.0 poor. Ligand efficiency normalises affinity by heavy-atom "
        "count and is the fairer comparison across molecules of different size.")
    doc.add_paragraph(
        "Limitations. These poses come from a single docking replica at the "
        "exhaustiveness recorded above, in the absence of explicit solvent. Docking "
        "scores are not free energies, and the ranking is more reliable than the "
        "absolute values. Treat this as a hypothesis-generating screen rather than "
        "a lead-identification result, and re-run anything promising at higher "
        "exhaustiveness with replicas before acting on it.")


def build_string_docx(r):
    """STRING interaction network report: network image + narrative + tables."""
    from docx import Document
    from docx.shared import Inches
    import pandas as pd

    names = ", ".join(r.get("input", [])) or "protein"
    doc = Document()
    doc.add_heading(f"MUMO Interaction Network Report — {names}", level=0)
    doc.add_paragraph("STRING protein–protein associations (known + predicted). "
                       "Combined score 0–1 from several evidence channels.")

    svg = r.get("network_svg")
    if svg:
        doc.add_heading("Interaction network", level=1)
        pw = browser = None
        try:
            pw, browser = new_browser()
            png = svg_to_png(svg, browser, width=900, height=700)
            doc.add_picture(io.BytesIO(png), width=Inches(6.0))
        except Exception as e:
            doc.add_paragraph(f"(Network image could not be rendered: {type(e).__name__}: {e})")
        finally:
            if browser:
                browser.close()
            if pw:
                pw.stop()

    # What each protein DOES — the layer STRING has no answer for. Built from
    # the same biology_blocks() the app panel renders, so the exported document
    # and the screen cannot disagree.
    try:
        from agents.string_deep import biology_blocks
        _order = list(r.get("input", [])) + [p.get("preferredName_B")
                                             for p in (r.get("partners") or [])]
        _blocks = biology_blocks(r.get("dossiers"), order=_order)
    except Exception:
        _blocks = []

    if _blocks:
        doc.add_heading("What these proteins do", level=1)
        doc.add_paragraph(
            "Curated biology from UniProt: molecular function, the reaction "
            "catalysed where there is one, the processes each protein takes part "
            "in, where in the cell it acts, and the diseases it is implicated in.")
        for b in _blocks:
            doc.add_heading(f'{b["gene"]} — {b["protein_name"]}', level=2)
            meta = " · ".join(x for x in (
                b.get("accession"),
                f'{b["length"]} aa' if b.get("length") else "") if x)
            if meta:
                doc.add_paragraph(meta)
            if b["summary"]:
                doc.add_paragraph(b["summary"])
            for label, text in b["rows"]:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(text)

    # Which of these proteins can actually be drugged. A network says what is
    # connected; this says which nodes are validated points of intervention.
    try:
        from agents.string_deep import druggability_map
        _doss = r.get("dossiers") or r.get("_dossiers") or {}
        _drug = r.get("_druggability") or (druggability_map(_doss) if _doss else {})
    except Exception:
        _drug = {}

    if _drug:
        doc.add_heading("Druggability", level=1)
        doc.add_paragraph(
            "Known modulators from ChEMBL, looked up by UniProt accession. "
            "A protein with approved drugs is a validated point of intervention; "
            "one with none is either genuinely hard to drug or simply unexplored.")
        _order = list(r.get("input", [])) + [p.get("preferredName_B")
                                             for p in (r.get("partners") or [])]
        _seen = set()
        _rows = []
        for g in _order:
            if g in _drug and g not in _seen:
                _seen.add(g)
                v = _drug[g]
                _rows.append({
                    "Protein": g,
                    "Druggability": v.get("verdict", ""),
                    "Modulators": str(v.get("n_mechanisms", 0)),
                    "Examples": ", ".join(x["name"] for x in v.get("drugs", [])[:3]) or "—",
                })
        if _rows:
            _add_df_table(doc, pd.DataFrame(_rows))

    narrative = r.get("narrative")
    if narrative:
        doc.add_heading("Report", level=1)
        _add_markdown(doc, narrative)

    partners = r.get("partners") or []
    if partners:
        doc.add_heading("Functional partners", level=1)
        rows = [{"Partner": p.get("preferredName_B", "?"),
                 "Score": round(p.get("score", 0), 3),
                 "Experimental": round(p.get("escore", 0), 3),
                 "Database": round(p.get("dscore", 0), 3),
                 "Text-mining": round(p.get("tscore", 0), 3)} for p in partners]
        _add_df_table(doc, pd.DataFrame(rows))

    enr = r.get("enrichment") or []
    if enr:
        doc.add_heading("Enriched pathways / functions", level=1)
        top = sorted(enr, key=lambda e: e.get("fdr", 1.0))[:12]
        rows = [{"Category": e.get("category", ""), "Term": e.get("description", ""),
                 "FDR": "{:.1e}".format(e.get("fdr", 1.0))} for e in top]
        _add_df_table(doc, pd.DataFrame(rows))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_ramachandran_docx(r):
    """Standalone backbone-geometry validation report for one structure."""
    from docx import Document
    from docx.shared import Inches
    import pandas as pd
    import ramachandran as ram

    res = r.get("result") or {}
    gene = res.get("gene") or r.get("target") or "target"

    doc = Document()
    doc.add_heading(f"MUMO Structure Validation — {gene}", level=0)
    if res.get("source"):
        doc.add_paragraph(f"Structure source: {res['source']}")
    if res.get("_error"):
        doc.add_paragraph(f"Validation could not run: {res['_error']}")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    doc.add_paragraph(ram.verdict(res))
    doc.add_paragraph(f"{res['n_scored']} of {res['n_residues']} residues scored "
                      f"(a chain's first and last residue have no phi/psi pair).")

    if r.get("narrative"):
        doc.add_heading("Report", level=1)
        _add_markdown(doc, r["narrative"])

    doc.add_heading("Ramachandran plot", level=1)
    svg = ram.plot_svg(res, title=f"Ramachandran plot — {gene}")
    if svg:
        pw = browser = None
        try:
            pw, browser = new_browser()
            png = svg_to_png(svg, browser, width=760, height=860)
            doc.add_picture(io.BytesIO(png), width=Inches(5.6))
        except Exception as e:
            doc.add_paragraph(f"(Plot could not be rendered: {type(e).__name__}: {e})")
        finally:
            try:
                if browser:
                    browser.close()
                if pw:
                    pw.stop()
            except Exception:
                pass

    doc.add_heading("Summary", level=1)
    _add_kv_table(doc, [
        ("Residues scored", res["n_scored"]),
        ("Favoured", f'{res["counts"]["favoured"]} ({res["pct"]["favoured"]}%)'),
        ("Allowed", f'{res["counts"]["allowed"]} ({res["pct"]["allowed"]}%)'),
        ("Outliers", f'{res["counts"]["outlier"]} ({res["pct"]["outlier"]}%)'),
    ])

    outs = res.get("outliers") or []
    if outs:
        doc.add_heading("Outlier residues", level=1)
        doc.add_paragraph(
            "These residues have backbone angles outside the regions normally "
            "accessible to their residue type. Check whether any lie in or near "
            "a binding site before trusting docking results for this structure.")
        _add_df_table(doc, pd.DataFrame({"Residue": outs}))

    doc.add_heading("Method and limitations", level=1)
    doc.add_paragraph(
        "Phi/psi torsion angles are computed directly from backbone N, CA and C "
        "coordinates using standard IUPAC geometry. Residues are scored against "
        "region definitions specific to their type — glycine reaches mirror "
        "regions no side-chain-bearing residue can, and proline's ring "
        "restricts phi. Torsions are not computed across chain breaks.")
    doc.add_paragraph(res.get("note", ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_metabolism_docx(r):
    """Metabolism report: narrative, the ROUTE DIAGRAM, then the table.

    The diagram is the point. A route written as "Aromatic hydroxylation ->
    O-glucuronidation" tells a chemist almost nothing on its own — they need to
    see which position was hydroxylated and what the conjugate looks like. So
    the structures are rendered and rasterized the same way the docking figures
    are, and the table follows as the precise reference.
    """
    from docx import Document
    from docx.shared import Inches
    import pandas as pd     # local: report_writer takes DataFrames in elsewhere

    pred = r.get("prediction") or {}
    mets = pred.get("metabolites") or []

    doc = Document()
    doc.add_heading(f"MUMO Metabolism Report — {r.get('lig_label', 'compound')}", level=0)
    doc.add_paragraph(f"Parent SMILES: {r.get('lig_smiles', '')}")
    n1 = sum(1 for m in mets if m.get("phase") == "I")
    n2 = sum(1 for m in mets if m.get("phase") == "II")
    doc.add_paragraph(f"{len(mets)} metabolites shown ({n1} phase I, {n2} phase II) "
                      f"of {pred.get('n_generated', 0)} generated.")

    if r.get("narrative"):
        doc.add_heading("Report", level=1)
        _add_markdown(doc, r["narrative"])

    # ── the route diagram ────────────────────────────────────────────────
    try:
        from agents.metabolism import pathway_svg
        svg = pathway_svg(pred, max_routes=6)
        if svg:
            pw = browser = None
            try:
                pw, browser = new_browser()
                png = svg_to_png(svg, browser, width=1500, height=1200)
                doc.add_heading("Predicted routes", level=1)
                doc.add_picture(io.BytesIO(png), width=Inches(6.4))
            finally:
                try:
                    if browser:
                        browser.close()
                    if pw:
                        pw.stop()
                except Exception:
                    pass
    except Exception as e:
        # the tables below still carry the substance — never lose the report
        doc.add_paragraph(f"[Route diagram unavailable: {type(e).__name__}: {e}]")

    if mets:
        doc.add_heading("Predicted metabolites", level=1)
        doc.add_paragraph("Score ranks likelihood; it is not a predicted amount.")
        rows = [{"Rank": i, "Metabolite (SMILES)": m["smiles"],
                 "Transformation": m["name"], "Phase": m["phase"],
                 "Score": m["score"], "Route": " -> ".join(m["pathway"])}
                for i, m in enumerate(mets, 1)]
        _add_df_table(doc, pd.DataFrame(rows))

    doc.add_heading("Method and limitations", level=1)
    doc.add_paragraph(pred.get("citation", ""))
    doc.add_paragraph(
        "Metabolites are generated by applying literature-derived phase I and "
        "phase II reaction rules to the parent structure, then to the products, "
        "so a phase I step can be followed by a phase II conjugation. Scores "
        "multiply along a route and rank relative likelihood only. These are "
        "computational hypotheses for experimental testing, not measurements, "
        "and absence from this list is not evidence a metabolite does not form.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_admet_docx(r):
    """ADMET / drug-likeness report: tables + beginner narrative (no images)."""
    from docx import Document

    doc = Document()
    doc.add_heading(f"MUMO ADMET Report — {r.get('lig_label', 'ligand')}", level=0)
    doc.add_paragraph(f"SMILES: {r.get('lig_smiles', '')}")

    qm = r.get("qm")
    if qm and not qm.get("_error"):
        doc.add_heading("Frontier molecular orbitals", level=1)
        doc.add_paragraph(f"Method: {qm.get('method', 'GFN2-xTB')}")
        _add_kv_table(doc, [
            ("HOMO", f"{qm['homo_ev']:.3f} eV"),
            ("LUMO", f"{qm['lumo_ev']:.3f} eV"),
            ("HOMO-LUMO gap", f"{qm['gap_ev']:.3f} eV"),
        ] + ([("Dipole moment", f"{qm['dipole_debye']:.2f} D")]
             if qm.get("dipole_debye") is not None else []))
        try:
            import viz_string as _vs
            _svg = _vs.orbital_svg(qm, dark=False)
            if _svg:
                pw = browser = None
                try:
                    pw, browser = new_browser()
                    png = svg_to_png(_svg, browser, width=700, height=430)
                    doc.add_picture(io.BytesIO(png), width=Inches(5.2))
                finally:
                    if browser:
                        browser.close()
                    if pw:
                        pw.stop()
        except Exception as e:
            doc.add_paragraph(f"(Orbital diagram could not be rendered: {e})")
        if qm.get("interpretation"):
            doc.add_paragraph(qm["interpretation"])

    narrative = r.get("narrative")
    if narrative:
        doc.add_heading("Report", level=1)
        _add_markdown(doc, narrative)

    dl = r.get("druglikeness") or {}
    if dl:
        doc.add_heading("Drug-likeness", level=1)
        _add_kv_table(doc, list(dl.items()))

    adm = r.get("admet_ml") or {}
    if adm and "_error" not in adm:
        doc.add_heading("ADMET-AI predictions", level=1)
        doc.add_paragraph("Pretrained ML models (Therapeutics Data Commons / Chemprop). "
                           "Classifier endpoints are probabilities 0–1; regression endpoints "
                           "are in native units.")
        _add_kv_table(doc, list(adm.items()))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# structure export: the raw docked geometry in standard formats, so a user can
# open the pose in Discovery Studio / Maestro / BIOVIA / PyMOL / ChimeraX etc.
# Everything is derived from the (persisted) complex PDB, so this works for both
# fresh and reloaded results without needing the original run's temp files.
# ─────────────────────────────────────────────────────────────────────────────

def _split_complex_pdb(pdb_text):
    """Split a docked complex PDB into (receptor_pdb, ligand_pdb). The docked
    ligand is the HETATM group named LIG; everything else is the receptor."""
    receptor, ligand = [], []
    for ln in pdb_text.splitlines():
        rec = ln[:6].strip()
        if rec == "ATOM":
            receptor.append(ln)
        elif rec == "HETATM":
            (ligand if ln[17:20].strip() == "LIG" else receptor).append(ln)
        elif rec in ("TER", "HEADER", "CRYST1", "SEQRES"):
            receptor.append(ln)
    rec_txt = ("\n".join(receptor) + "\nEND\n") if receptor else ""
    lig_txt = ("\n".join(ligand) + "\nEND\n") if ligand else ""
    return rec_txt, lig_txt


def _corrected_ligand_mol(lig_pdb, smiles=None):
    """RDKit ligand mol from the docked PDB block with bond orders restored from the
    known SMILES (a docked PDB has coordinates but NO bonds). Returns None on failure."""
    try:
        from rdkit import Chem
        from rdkit.Chem import AllChem
    except Exception:
        return None
    mol = (Chem.MolFromPDBBlock(lig_pdb, sanitize=True, removeHs=False)
           or Chem.MolFromPDBBlock(lig_pdb, sanitize=False, removeHs=False))
    if mol is None:
        return None
    if smiles:
        try:
            tmpl = Chem.MolFromSmiles(smiles)
            if tmpl is not None:
                mol = AllChem.AssignBondOrdersFromTemplate(tmpl, mol)
        except Exception:
            pass
    return mol


def _add_ligand_conect(complex_pdb, lig_mol):
    """Append CONECT records for the ligand's bonds to a complex PDB, so viewers
    draw the ligand with its REAL connectivity instead of guessing bonds from atom
    distances (which tangles folded/large ligands). No-op if the atom counts don't
    line up."""
    if lig_mol is None:
        return complex_pdb
    serials = []
    for ln in complex_pdb.splitlines():
        if ln.startswith("HETATM") and ln[17:20].strip() == "LIG":
            try:
                serials.append(int(ln[6:11]))
            except ValueError:
                return complex_pdb
    if lig_mol.GetNumAtoms() != len(serials):
        return complex_pdb
    conect = [f"CONECT{serials[b.GetBeginAtomIdx()]:>5}{serials[b.GetEndAtomIdx()]:>5}"
              for b in lig_mol.GetBonds()]
    lines = [ln for ln in complex_pdb.splitlines() if ln.strip() != "END"]
    return "\n".join(lines + conect + ["END"]) + "\n"


def _ligand_sdf(lig_pdb, smiles=None):
    """Ligand PDB block → MDL molblock (.sdf/.mol) text. Uses the known SMILES
    to restore correct bond orders (PDB has none). Returns None on failure."""
    try:
        from rdkit import Chem
        from rdkit.Chem import AllChem
    except Exception:
        return None
    mol = Chem.MolFromPDBBlock(lig_pdb, sanitize=True, removeHs=False)
    if mol is None:
        mol = Chem.MolFromPDBBlock(lig_pdb, sanitize=False, removeHs=False)
    if mol is None:
        return None
    if smiles:
        try:
            tmpl = Chem.MolFromSmiles(smiles)
            if tmpl is not None:
                mol = AllChem.AssignBondOrdersFromTemplate(tmpl, mol)
        except Exception:
            pass  # keep the perceived-connectivity mol if template matching fails
    try:
        # MolToMolBlock ends at "M  END" (a .mol block); append the SDF record
        # terminator so the .sdf file is valid for strict parsers.
        return Chem.MolToMolBlock(mol).rstrip() + "\n$$$$\n"
    except Exception:
        return None


def _sybyl_type(atom):
    """Minimal SYBYL atom type for a MOL2 file (enough for viewers to read)."""
    from rdkit import Chem
    sym = atom.GetSymbol()
    hyb = atom.GetHybridization()
    if atom.GetIsAromatic() and sym in ("C", "N"):
        return f"{sym}.ar"
    if sym == "C":
        return {Chem.HybridizationType.SP: "C.1", Chem.HybridizationType.SP2: "C.2"}.get(hyb, "C.3")
    if sym == "N":
        return {Chem.HybridizationType.SP: "N.1", Chem.HybridizationType.SP2: "N.2"}.get(hyb, "N.3")
    if sym == "O":
        return "O.2" if hyb == Chem.HybridizationType.SP2 else "O.3"
    if sym == "S":
        return "S.3"
    if sym == "P":
        return "P.3"
    return sym


def _ligand_mol2(lig_pdb, smiles=None):
    """Ligand PDB block → MOL2 (TRIPOS) text via a small RDKit writer — no
    OpenBabel. Bond orders come from `smiles` when given. Returns None on failure."""
    try:
        from rdkit import Chem
        from rdkit.Chem import AllChem
    except Exception:
        return None
    mol = Chem.MolFromPDBBlock(lig_pdb, sanitize=True, removeHs=False)
    if mol is None:
        mol = Chem.MolFromPDBBlock(lig_pdb, sanitize=False, removeHs=False)
    if mol is None:
        return None
    if smiles:
        try:
            tmpl = Chem.MolFromSmiles(smiles)
            if tmpl is not None:
                mol = AllChem.AssignBondOrdersFromTemplate(tmpl, mol)
        except Exception:
            pass
    try:
        AllChem.ComputeGasteigerCharges(mol)
    except Exception:
        pass
    try:
        conf = mol.GetConformer()
        atom_lines = []
        for i, atom in enumerate(mol.GetAtoms()):
            p = conf.GetAtomPosition(i)
            try:
                q = float(atom.GetProp("_GasteigerCharge"))
                if q != q:      # NaN guard
                    q = 0.0
            except Exception:
                q = 0.0
            atom_lines.append(
                f"{i+1:>7} {atom.GetSymbol()+str(i+1):<8}{p.x:10.4f}{p.y:10.4f}{p.z:10.4f} "
                f"{_sybyl_type(atom):<6}{1:>4} LIG {q:>10.4f}")
        BOND = {Chem.BondType.SINGLE: "1", Chem.BondType.DOUBLE: "2",
                Chem.BondType.TRIPLE: "3", Chem.BondType.AROMATIC: "ar"}
        bond_lines = [f"{j+1:>6} {b.GetBeginAtomIdx()+1:>5} {b.GetEndAtomIdx()+1:>5} "
                      f"{BOND.get(b.GetBondType(), '1'):>2}"
                      for j, b in enumerate(mol.GetBonds())]
        return "\n".join([
            "@<TRIPOS>MOLECULE", "LIG",
            f" {mol.GetNumAtoms()} {mol.GetNumBonds()} 0 0 0", "SMALL", "GASTEIGER", "",
            "@<TRIPOS>ATOM", *atom_lines,
            "@<TRIPOS>BOND", *bond_lines, ""])
    except Exception:
        return None


def build_structure_zip(r):
    """Bundle the docked structures for external viewers. Per ligand: the docked
    complex, ligand, and receptor as PDB, plus ligand SDF (correct bonds) and
    MOL2 (best-effort). Returns zip bytes, or None if there's nothing to export."""
    import os
    import zipfile

    viz = r.get("viz") or {}
    rdf = r.get("rdf")
    meta = r.get("meta") or {}
    gene = re.sub(r"[^A-Za-z0-9_.-]", "_", str(meta.get("gene", "target"))) or "target"

    smiles_by = {}
    rank_by = {}
    if rdf is not None:
        for idx, row in rdf.iterrows():
            # A multi-target run keys its viz by "TARGET · ligand" so the same
            # ligand docked against two proteins doesn't overwrite itself, so
            # register BOTH forms — the plain name for a single-target run and
            # the composite for a screen.
            keys = [row.get("Ligand")]
            if row.get("Target"):
                keys.append(f'{row["Target"]} · {row.get("Ligand")}')
            for k in keys:
                smiles_by[k] = row.get("SMILES")
                rank_by[k] = idx                  # rdf is already sorted best→worst

    def _label_parts(label):
        """(target_or_None, ligand_name) from a viz key.

        A single-target run keys viz by plain ligand name. A multi-target run
        keys it "TARGET · ligand" instead (see pipeline_core.run_job) — parsing
        that same key back apart is what lets every use below (filenames, the
        receptor file, the README) name the correct target per ligand, instead
        of just inheriting whatever the raw dict key happened to look like.
        """
        s = str(label)
        if " · " in s:
            t, _, lig = s.partition(" · ")
            return t, lig
        return None, s

    used_stems = set()

    def _stem(label, fallback_rank):
        """A unique, rank-prefixed filename stem for one ligand.

        Compound names are truncated to keep paths sane, and GC-MS peak lists are
        full of homologs that differ only in a suffix ('…-, methyl ester' vs
        '…-, octadecyl ester') — those collide once truncated, and a zip member
        written twice means the second silently replaces the first on extraction,
        losing a ligand's structures. The rank prefix makes the name unique AND
        lets a file be matched to its row in the results table.

        A multi-target run's ligand name is prefixed with ITS target
        ("PLA2-lupeol") rather than sanitizing the raw "PLA2 · lupeol" key
        verbatim — the raw key's middle dot and spaces would otherwise turn
        into a run of underscores ("PLA2___lupeol") that is valid but reads as
        a rendering glitch, not a deliberate name.
        """
        target, lig = _label_parts(label)
        base_src = f"{target}-{lig}" if target else lig
        base = re.sub(r"[^A-Za-z0-9_.-]", "_", str(base_src))[:40] or "ligand"
        stem = f"{rank_by.get(label, fallback_rank):02d}_{base}"
        if stem in used_stems:                    # belt and braces
            n = 2
            while f"{stem}_{n}" in used_stems:
                n += 1
            stem = f"{stem}_{n}"
        used_stems.add(stem)
        return stem

    buf = io.BytesIO()
    receptors_written = set()   # per TARGET, not once per zip — see below
    wrote_any = False
    targets_list = meta.get("targets") or ([meta["gene"]] if meta.get("gene") else [])
    readme = ["MUMO — docked structures", "=" * 25, "",
              f"Target(s): {', '.join(targets_list) or meta.get('gene', '?')}",
              f"Pocket: {meta.get('pocket', 'n/a')}", "",
              "Files per ligand:",
              "  *_complex.pdb  — receptor + docked ligand pose (open this to see the pose)",
              "  *_ligand.pdb   — docked ligand only",
              "  *_ligand.sdf   — docked ligand with correct bond orders",
              "  *_ligand.mol2  — docked ligand (if available)",
              "  <target>_receptor.pdb — that target's protein only "
              "(one per target — do NOT mix a receptor with a ligand docked "
              "against a DIFFERENT target)", "",
              "Open in Discovery Studio, Maestro/BIOLuminate, BIOVIA, PyMOL, ChimeraX, etc.",
              "Note: structures reflect the exact docked pose from this run.", "",
              "Ligands:"]

    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        for i, (label, entry) in enumerate(viz.items(), 1):
            target, lig_name = _label_parts(label)
            safe = _stem(label, i)
            try:
                with open(entry["complex"]) as f:
                    complex_pdb = f.read()
            except Exception:
                continue
            rec_pdb, lig_pdb = _split_complex_pdb(complex_pdb)
            smi = smiles_by.get(label)
            # correct-bond ligand mol → used to add CONECT records so viewers show
            # the ligand with real bonds (not distance-guessed tangles)
            lig_mol = _corrected_ligand_mol(lig_pdb, smi) if lig_pdb else None
            # `gene` (the whole screen's joined label) prefixes the filename on
            # a single-target run, same as always. On a multi-target run it
            # would ONLY duplicate what `safe` already says — "PLA2-lupeol"
            # under an outer "PLA2___2PE4_" prefix reads like a rendering
            # glitch, not a name — so skip it there; `safe` alone already
            # names the target unambiguously.
            complex_name = f"{safe}_complex.pdb" if target else f"{gene}_{safe}_complex.pdb"
            z.writestr(complex_name, _add_ligand_conect(complex_pdb, lig_mol))
            wrote_any = True
            if lig_pdb:
                # ligand.pdb from the corrected mol (RDKit writes CONECT records) so
                # the standalone ligand also opens with correct bonds; else raw block
                lig_pdb_out = lig_pdb
                if lig_mol is not None:
                    try:
                        from rdkit import Chem
                        lig_pdb_out = Chem.MolToPDBBlock(lig_mol)
                    except Exception:
                        pass
                z.writestr(f"{safe}_ligand.pdb", lig_pdb_out)
                sdf = _ligand_sdf(lig_pdb, smi)
                if sdf:
                    z.writestr(f"{safe}_ligand.sdf", sdf)
                mol2 = _ligand_mol2(lig_pdb, smi)
                if mol2:
                    z.writestr(f"{safe}_ligand.mol2", mol2)
            # ONE receptor file per TARGET, not once for the whole zip. The old
            # "write it only the first time" guard meant a multi-target screen
            # silently exported only the FIRST target's receptor — pairing a
            # second-target ligand with that leftover file in an external
            # viewer would show wrong, mismatched interactions with no error
            # anywhere to explain why.
            target_slug = re.sub(r"[^A-Za-z0-9_.-]", "_", str(target or gene)) or "target"
            if rec_pdb and target_slug not in receptors_written:
                z.writestr(f"{target_slug}_receptor.pdb", rec_pdb)
                receptors_written.add(target_slug)
            readme.append(f"  - {lig_name} — vs {target}" if target else f"  - {lig_name}")
        z.writestr("README.txt", "\n".join(readme) + "\n")

    return buf.getvalue() if wrote_any else None
